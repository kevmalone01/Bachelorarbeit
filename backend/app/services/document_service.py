import os
import tempfile
import logging
import platform
import subprocess
import shutil
from typing import Dict, Any, List, Optional, BinaryIO, Tuple
import base64

import docx2pdf

from flask import current_app

from werkzeug.utils import secure_filename

from app.db import db
from app.models.document import Document
from app.models.placeholder import Placeholder

logger = logging.getLogger(__name__)

class DocumentService:
    """Service for document operations including preview and processing"""
    
    def __init__(self):
        self.allowed_extensions = {'pdf', 'docx', 'doc'}
        self.allowed_mime_types = {
            'application/pdf',
            'application/msword',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/x-msword',
            'application/octet-stream',
            'application/msword; charset=binary',
            'application/vnd.ms-word',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document; charset=binary'
        }
        self.upload_folder = None
        logger.info("DocumentService initialized with allowed extensions: %s", self.allowed_extensions)
        logger.info("DocumentService initialized with allowed MIME types: %s", self.allowed_mime_types)    
    
    def _get_upload_folder(self) -> str:
        """Get upload folder from app config or default"""
        if not self.upload_folder:
            try:
                folder = current_app.config.get('UPLOAD_FOLDER', 'uploads')
                logger.debug("Upload folder from config: %s", folder)
                return folder
            except RuntimeError:
                folder = os.getenv("UPLOAD_FOLDER", "uploads")
                logger.debug("Upload folder from environment: %s", folder)
                return folder
        logger.debug("Using cached upload folder: %s", self.upload_folder)
        return self.upload_folder
    
    def _allowed_file(self, filename: str, content_type: str = None) -> bool:
        """Check if file extension and MIME type are allowed"""
        # Check file extension
        extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
        extension_allowed = extension in self.allowed_extensions
        
        # Log initial validation info
        logger.info("Validating file: %s (type: %s)", filename, content_type)
        logger.debug("File extension: %s, Allowed extensions: %s", extension, self.allowed_extensions)
        
        # Für DOCX: Immer erlauben, wenn die Endung stimmt
        if extension == 'docx':
            logger.info("DOCX-Datei erkannt, Endung stimmt. Upload wird erlaubt.")
            return True
        
        # Für PDF und DOC: Extension und MIME-Type prüfen
        mime_type_allowed = True
        if content_type:
            normalized_content_type = content_type.lower().split(';')[0].strip()
            mime_type_allowed = any(
                allowed_type.lower().split(';')[0].strip() == normalized_content_type
                for allowed_type in self.allowed_mime_types
            )
            logger.debug("Normalized content type: %s", normalized_content_type)
            logger.debug("Allowed MIME types: %s", self.allowed_mime_types)
        
        is_allowed = extension_allowed or mime_type_allowed
        
        if not is_allowed:
            logger.warning("File validation failed:")
            logger.warning("- Filename: %s", filename)
            logger.warning("- Content type: %s", content_type)
            logger.warning("- Extension: %s", extension)
            logger.warning("- Extension allowed: %s", extension_allowed)
            logger.warning("- MIME type allowed: %s", mime_type_allowed)
            logger.warning("- Allowed extensions: %s", self.allowed_extensions)
            logger.warning("- Allowed MIME types: %s", self.allowed_mime_types)
            raise ValueError(f"Nur PDF und DOCX Dateien werden unterstützt. Ihre Datei hat die Endung '{extension}' und den MIME-Type '{content_type}'.")
        else:
            logger.info("File validation successful")
            logger.debug("- Extension allowed: %s", extension_allowed)
            logger.debug("- MIME type allowed: %s", mime_type_allowed)
        
        return is_allowed
    
    def save_document(self, file: BinaryIO, filename: str, content_type: str = None) -> str:
        """Save uploaded document and return the path"""
        logger.info("Saving document: %s (type: %s)", filename, content_type)
        
        # If content type is not provided, try to determine it from the filename
        if not content_type:
            extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
            if extension == 'docx':
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            elif extension == 'pdf':
                content_type = 'application/pdf'
            logger.info("Determined content type from extension: %s", content_type)
        
        if not self._allowed_file(filename, content_type):
            logger.error("File type not allowed for file: %s (type: %s)", filename, content_type)
            raise ValueError(f"File type not allowed. Allowed types: {', '.join(self.allowed_extensions)}")
        
        upload_folder = self._get_upload_folder()
        os.makedirs(upload_folder, exist_ok=True)
        logger.debug("Created upload folder: %s", upload_folder)
        
        secure_name = secure_filename(filename)
        file_path = os.path.join(upload_folder, secure_name)
        logger.debug("Saving file to: %s", file_path)
        
        with open(file_path, 'wb') as f:
            f.write(file.read())
        
        logger.info("Successfully saved document to: %s", file_path)
        return file_path
    
    def create_document_preview(self, document_id: int, placeholder_values: Dict[str, Any], content_html: Optional[str] = None, export_format: str = 'docx') -> Dict[str, Any]:
        """Create a preview of a document with placeholders filled in
        
        Args:
            document_id: ID of the document template
            placeholder_values: Dictionary of placeholder values keyed by placeholder name
            content_html: Optional HTML content to use instead of loading from file
            export_format: Format to export ('pdf' or 'docx')
            
        Returns:
            Dictionary with preview info including base64 data and mime type
        """
        logger.info("Creating document preview for document ID: %d, format: %s", document_id, export_format)
        logger.debug("Placeholder values: %s", placeholder_values)
        if content_html:
            logger.info("Using provided contentHtml (length: %d) instead of file", len(content_html))
        
        document = Document.query.get(document_id)
        if not document:
            logger.error("Document with ID %d not found", document_id)
            raise ValueError(f"Document with ID {document_id} not found")
        
        # If document has no file_path and is not a template, try to find the template it was created from
        if not document.file_path and not document.is_template:
            logger.warning("Document with ID %d has no file_path, searching for template", document_id)
            # Try to find template by matching title (documents created from templates often have the same title)
            template = Document.query.filter_by(
                title=document.title,
                is_template=True
            ).first()
            
            if template and template.file_path:
                # Check if the file actually exists before using it
                import os
                if os.path.exists(template.file_path):
                    logger.info("Found template by title for document %d: Template ID %d, file_path: %s", 
                              document_id, template.id, template.file_path)
                    # Temporarily set file_path for processing
                    document.file_path = template.file_path
                else:
                    logger.warning("Template file_path does not exist: %s, searching for other templates", template.file_path)
                    template = None  # Reset to search for other templates
            
            if not template or not document.file_path:
                # Try to find any template with a valid file_path as fallback
                templates = Document.query.filter_by(is_template=True).filter(
                    Document.file_path.isnot(None)
                ).all()
                
                # Find the first template with an existing file
                import os
                for t in templates:
                    if t.file_path and os.path.exists(t.file_path):
                        logger.info("Using template (ID %d) with valid file_path for document %d: %s", 
                                 t.id, document_id, t.file_path)
                        document.file_path = t.file_path
                        break
                else:
                    # No template with valid file found
                    logger.error("Document with ID %d has no file_path and no template with valid file found", document_id)
                    raise ValueError(f"Document with ID {document_id} has no file_path and cannot find template with valid file")
        
        if not document.file_path:
            logger.error("Document with ID %d has no file_path", document_id)
            raise ValueError(f"Document with ID {document_id} has no file_path")
        
        # Verify the file actually exists
        import os
        if not os.path.exists(document.file_path):
            logger.error("Document file_path does not exist: %s for document ID %d", document.file_path, document_id)
            # If file doesn't exist, try to find the most recent template with a valid file
            if not document.is_template:
                logger.info("Searching for most recent template with valid file_path for document %d", document_id)
                templates = Document.query.filter_by(is_template=True).filter(
                    Document.file_path.isnot(None)
                ).order_by(Document.updated_at.desc(), Document.created_at.desc()).all()
                
                for t in templates:
                    if t.file_path and os.path.exists(t.file_path):
                        logger.info("Found valid template (ID %d) with file_path: %s, using for document %d", 
                                 t.id, t.file_path, document_id)
                        document.file_path = t.file_path
                        break
                else:
                    raise ValueError(f"Document file not found at path: {document.file_path} and no valid template found")
            else:
                raise ValueError(f"Document file not found at path: {document.file_path}")
        
        # Get associated placeholders from the document's placeholders JSON field
        placeholders = document.placeholders or []
        logger.debug("Found %d placeholders for document", len(placeholders))
        
        # Validate required placeholders are provided
        missing_required = []
        for placeholder in placeholders:
            if placeholder.get('required', False) and placeholder.get('name') not in placeholder_values:
                missing_required.append(placeholder.get('name'))
        
        if missing_required:
            logger.error("Missing required placeholders: %s", ", ".join(missing_required))
            raise ValueError(f"Missing required placeholders: {', '.join(missing_required)}")
        
        # Determine document type from file_path if document_type is not a MIME type
        # If document_type is a custom type (like "Mandantenonboarding"), use file_path extension
        actual_document_type = document.document_type
        if document.file_path:
            file_ext = os.path.splitext(document.file_path)[1].lower()
            if file_ext == '.pdf':
                actual_document_type = 'application/pdf'
            elif file_ext in ['.docx', '.doc']:
                actual_document_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            logger.debug("Determined document type from file extension: %s -> %s", file_ext, actual_document_type)
        
        # If content_html is provided, check if we should use original file or HTML
        # IMPORTANT: Prefer original DOCX file to preserve formatting
        # Only use HTML conversion if no DOCX file is available
        if content_html:
            # Check if we have a DOCX file available
            if document.file_path and document.file_path.lower().endswith(('.docx', '.doc')):
                logger.info("Document %d has DOCX file_path, using original file to preserve formatting (ignoring contentHtml for better formatting)", document_id)
                # Use the original DOCX file processing - this preserves all formatting
                # The content_html is ignored in favor of the original file for better formatting
                logger.info("Processing DOCX document: %s, export format: %s", document.title, export_format)
                return self._process_docx_preview(document, placeholder_values, export_format=export_format)
            elif document.file_path and document.file_path.lower().endswith('.pdf'):
                logger.info("Document %d has PDF file_path, using original file", document_id)
                return self._process_pdf_preview(document, placeholder_values)
            else:
                # No DOCX/PDF file available, use HTML conversion
                logger.info("No DOCX/PDF file available for document %d, using HTML conversion to %s", document_id, export_format.upper())
                try:
                    return self._process_html_content(document, content_html, placeholder_values, export_format=export_format)
                except Exception as e:
                    logger.error("Error processing HTML content: %s", str(e), exc_info=True)
                    # Re-raise to be handled by the calling function
                    raise
        
        # Process document based on type
        if actual_document_type == 'application/pdf' or (document.file_path and document.file_path.lower().endswith('.pdf')):
            logger.info("Processing PDF document: %s", document.title)
            return self._process_pdf_preview(document, placeholder_values)
        elif actual_document_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or (document.file_path and document.file_path.lower().endswith(('.docx', '.doc'))):
            logger.info("Processing DOCX document: %s, export format: %s", document.title, export_format)
            return self._process_docx_preview(document, placeholder_values, export_format=export_format)
        else:
            logger.error("Unsupported document type: %s (file_path: %s)", document.document_type, document.file_path)
            raise ValueError(f"Unsupported document type: {document.document_type}. File path: {document.file_path}")
    
    def _process_pdf_preview(self, document: Document, placeholder_values: Dict[str, Any]) -> Dict[str, Any]:
        """Process PDF document with placeholders
        
        Note: This is a simplified implementation. In a real-world scenario,
        you might use libraries like PyPDF2, pdfrw, or ReportLab to modify PDFs.
        """
        try:
            logger.debug("Processing PDF preview for document: %s", document.title)
            # In a real implementation, you would manipulate the PDF here
            # For now, we'll just read the original file
            with open(document.file_path, 'rb') as f:
                file_data = f.read()
            
            # Return base64 encoded data
            base64_data = base64.b64encode(file_data).decode('utf-8')
            logger.info("Successfully created PDF preview for: %s", document.title)
            return {
                'preview_data': base64_data,
                'mime_type': 'application/pdf',
                'filename': f"{document.title}_preview.pdf"
            }
        except Exception as e:
            logger.error("Error processing PDF preview: %s", str(e), exc_info=True)
            current_app.logger.error(f"Error processing PDF preview: {str(e)}")
            raise Exception(f"Failed to create PDF preview: {str(e)}")
    
    def _convert_with_docx2pdf(self, docx_path: str, output_path: str) -> Tuple[bool, str]:
        """
        Convert DOCX to PDF using docx2pdf library
        
        Args:
            docx_path: Path to the DOCX file
            output_path: Path where the PDF should be saved
            
        Returns:
            Tuple of (success, error_message)
        """
        try:
            from docx2pdf import convert
            
            logger.info("Converting DOCX to PDF using docx2pdf library. Input: %s, Output: %s", docx_path, output_path)
            output_dir = os.path.dirname(output_path)
            
            # Ensure output directory exists
            os.makedirs(output_dir, exist_ok=True)
            
            # On macOS, docx2pdf requires LibreOffice or Microsoft Word
            # Check if we're on macOS and warn if needed
            if platform.system() == 'Darwin':
                logger.info("Running on macOS - docx2pdf may require LibreOffice or Microsoft Word")
            
            # Initialize COM for Windows if needed
            if platform.system() == 'Windows':
                try:
                    import pythoncom
                    pythoncom.CoInitialize()
                    logger.debug("COM initialized for Windows")
                except Exception as com_error:
                    logger.warning("Failed to initialize COM: %s", str(com_error))
            
            # Convert the document with timeout handling
            logger.info("Starting DOCX to PDF conversion...")
            try:
                convert(docx_path, output_path)
                logger.info("Conversion command completed")
            except Exception as convert_error:
                logger.error("Conversion command failed: %s", str(convert_error), exc_info=True)
                return False, f"Conversion failed: {str(convert_error)}"
            
            # Check if the conversion was successful
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                logger.info("Successfully converted DOCX to PDF using docx2pdf. File size: %d bytes", file_size)
                return True, ""
            else:
                logger.error("docx2pdf did not create output file at: %s", output_path)
                return False, "PDF file not created by docx2pdf"
                
        except Exception as e:
            logger.error("Error using docx2pdf for conversion: %s", str(e), exc_info=True)
            return False, str(e)
        finally:
            # Uninitialize COM if we're on Windows
            if platform.system() == 'Windows':
                try:
                    import pythoncom
                    pythoncom.CoUninitialize()
                    logger.debug("COM uninitialized for Windows")
                except Exception:
                    pass
    
    def _convert_with_pypdf_docx2pdf(self, docx_path: str) -> Tuple[bool, str, str]:
        """
        Convert DOCX to PDF using pypdf-docx2pdf library
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Tuple of (success, pdf_path, error_message)
        """
        # Check if pypdf-docx2pdf is installed
        try:
            import importlib.util
            spec = importlib.util.find_spec('pypdf_docx2pdf')
            if spec is None:
                logger.info("pypdf-docx2pdf module is not installed, skipping this conversion method")
                return False, "", "Module pypdf-docx2pdf is not installed"
        except ImportError:
            logger.info("Unable to check for pypdf-docx2pdf module, skipping this conversion method")
            return False, "", "Unable to check for pypdf-docx2pdf module"
            
        # Module is available, try to use it
        try:
            from pypdf_docx2pdf import convert as pdf_convert
            
            pdf_path = os.path.splitext(docx_path)[0] + '.pdf'
            logger.debug("Converting DOCX to PDF using pypdf-docx2pdf library")
            
            # Convert the document
            pdf_convert(docx_path, pdf_path)
            
            # Check if the conversion was successful
            if os.path.exists(pdf_path):
                logger.info("Successfully converted DOCX to PDF using pypdf-docx2pdf")
                return True, pdf_path, ""
            else:
                logger.error("pypdf-docx2pdf did not create output file at: %s", pdf_path)
                return False, "", "PDF file not created by pypdf-docx2pdf"
                
        except Exception as e:
            logger.error("Error using pypdf-docx2pdf for conversion: %s", str(e), exc_info=True)
            return False, "", str(e)
    
    def _try_convert_docx_to_pdf(self, docx_path: str) -> Tuple[bool, str, str]:
        """
        Try to convert DOCX to PDF using available methods
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Tuple of (success, pdf_path, error_message)
        """
        docx_dir = os.path.dirname(docx_path)
        pdf_path = os.path.splitext(docx_path)[0] + '.pdf'
        
        # Try using docx2pdf library first
        try:
            success, error_msg = self._convert_with_docx2pdf(docx_path, pdf_path)
            if success:
                return True, pdf_path, ""
            logger.warning("docx2pdf conversion failed: %s. Trying next method.", error_msg)
        except Exception as e:
            logger.warning("docx2pdf conversion error: %s. Trying next method.", str(e))
        
        # Try using pypdf-docx2pdf as a second option
        try:
            success, pypdf_path, error_msg = self._convert_with_pypdf_docx2pdf(docx_path)
            if success:
                return True, pypdf_path, ""
            logger.warning("pypdf-docx2pdf conversion failed: %s. Trying next method.", error_msg)
        except Exception as e:
            logger.warning("pypdf-docx2pdf conversion error: %s. Trying next method.", str(e))
            
        # If all conversion methods fail, return a failure result
        # This will cause the system to fall back to showing the DOCX file directly
        logger.info("All PDF conversion methods failed, falling back to DOCX format")
        return False, "", "All PDF conversion methods failed"
    
    def _process_docx_preview(self, document: Document, placeholder_values: Dict[str, Any], export_format: str = 'docx') -> Dict[str, Any]:
        """Process DOCX document with placeholders
        
        First replaces placeholders in the document, then converts to PDF if requested
        """
        try:
            import docx
            from docx.shared import Pt
            
            logger.debug("Processing DOCX preview for document: %s", document.title)
            
            # Open the template document
            logger.debug("Opening template document: %s", document.file_path)
            
            # Verify file exists before trying to open it
            import os
            if not os.path.exists(document.file_path):
                logger.error("Document file does not exist: %s", document.file_path)
                raise FileNotFoundError(f"Document file not found at path: {document.file_path}")
            
            if not os.access(document.file_path, os.R_OK):
                logger.error("Document file is not readable: %s", document.file_path)
                raise PermissionError(f"Document file is not readable: {document.file_path}")
            
            try:
                doc = docx.Document(document.file_path)
            except Exception as e:
                logger.error("Error opening DOCX file %s: %s", document.file_path, str(e), exc_info=True)
                raise Exception(f"Failed to open DOCX file: {str(e)}")
            
            # Replace placeholders in all paragraphs
            # IMPORTANT: DOCX files can have placeholders split across multiple runs
            # We need to check both individual runs and the full paragraph text
            placeholder_replacements = 0
            logger.info("Available placeholder values for replacement: %s", list(placeholder_values.keys()))
            logger.info("Placeholder values: %s", {k: str(v)[:50] for k, v in placeholder_values.items()})
            
            # First, scan all paragraphs to find what placeholders exist in the document
            all_paragraph_texts = [p.text for p in doc.paragraphs]
            all_text = ' '.join(all_paragraph_texts)
            import re
            found_placeholders = re.findall(r'\{\{([^}]+)\}\}', all_text)
            logger.info("Found placeholders in document: %s", list(set(found_placeholders)))
            
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text
                original_text = paragraph_text
                
                # Replace ALL placeholders in the paragraph before rebuilding
                for key, value in placeholder_values.items():
                    # Skip empty values
                    if value is None or value == '':
                        continue
                    
                    # Try multiple variations of the placeholder
                    variations = [
                        '{{' + key + '}}',           # {{key}}
                        '{{ ' + key + ' }}',         # {{ key }}
                        '{{' + key.strip() + '}}',  # {{key}} (trimmed)
                        '{{ ' + key.strip() + ' }}', # {{ key }} (trimmed)
                    ]
                    
                    # Also try case-insensitive variations
                    key_lower = key.lower()
                    if key_lower != key:
                        variations.extend([
                            '{{' + key_lower + '}}',
                            '{{ ' + key_lower + ' }}',
                        ])
                    
                    for placeholder in variations:
                        if placeholder in paragraph_text:
                            logger.debug("Replacing placeholder '%s' (variation: '%s') with value: '%s' in paragraph", key, placeholder, str(value))
                            paragraph_text = paragraph_text.replace(placeholder, str(value))
                            placeholder_replacements += 1
                            break  # Only replace once per key
                
                # Only rebuild paragraph if text changed
                if paragraph_text != original_text:
                    logger.debug("Paragraph text changed, rebuilding. Original: '%s', New: '%s'", original_text[:50], paragraph_text[:50])
                    # Clear paragraph and rebuild with replaced text
                    paragraph.clear()
                    if paragraph_text:  # Only add run if there's text
                        paragraph.add_run(paragraph_text)
            
            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text
                        original_cell_text = cell_text
                        
                        for key, value in placeholder_values.items():
                            # Skip empty values
                            if value is None or value == '':
                                continue
                            
                            # Try multiple variations of the placeholder
                            variations = [
                                '{{' + key + '}}',           # {{key}}
                                '{{ ' + key + ' }}',         # {{ key }}
                                '{{' + key.strip() + '}}',  # {{key}} (trimmed)
                                '{{ ' + key.strip() + ' }}', # {{ key }} (trimmed)
                            ]
                            
                            # Also try case-insensitive variations
                            key_lower = key.lower()
                            if key_lower != key:
                                variations.extend([
                                    '{{' + key_lower + '}}',
                                    '{{ ' + key_lower + ' }}',
                                ])
                            
                            for placeholder in variations:
                                if placeholder in cell_text:
                                    logger.debug("Replacing placeholder '%s' (variation: '%s') with value: '%s' in table cell", key, placeholder, str(value))
                                    cell_text = cell_text.replace(placeholder, str(value))
                                    placeholder_replacements += 1
                                    break  # Only replace once per key
                        
                        # Update cell text if changed
                        if cell_text != original_cell_text:
                            cell.text = cell_text
            
            logger.info("Replaced %d placeholders in document", placeholder_replacements)
            
            # Save to temporary file with placeholders replaced
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
                doc.save(tmp_docx.name)
                tmp_docx_path = tmp_docx.name
                logger.debug("Saved document with replacements to temp file: %s", tmp_docx_path)
            
            # Convert to requested format
            logger.info("Checking export format: export_format='%s', type=%s, lower()='%s', comparison result: %s", 
                       export_format, type(export_format).__name__, export_format.lower() if export_format else 'None', export_format.lower() == 'pdf' if export_format else False)
            if export_format and str(export_format).lower() == 'pdf':
                logger.info("Attempting to convert DOCX to PDF for document: %s", document.title)
                pdf_success = False
                pdf_data = None
                
                # Try converting DOCX to HTML first, then HTML to PDF (works on all platforms)
                try:
                    import mammoth
                    from reportlab.lib.pagesizes import A4
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    from reportlab.lib.styles import getSampleStyleSheet
                    from reportlab.lib.units import inch
                    import io
                    import html
                    import re
                    
                    logger.info("Trying DOCX -> HTML -> PDF conversion using mammoth and reportlab")
                    
                    # Convert DOCX to HTML using mammoth
                    # mammoth.convert_to_html() can accept either a file path (string) or a file-like object
                    # We'll pass the file path directly, which is the simplest approach
                    result = mammoth.convert_to_html(tmp_docx_path)
                    html_content = result.value
                    logger.info("Successfully converted DOCX to HTML, length: %d", len(html_content))
                    
                    # Log any warnings from mammoth conversion
                    if result.messages:
                        logger.warning("Mammoth conversion warnings: %s", result.messages)
                    
                    # Convert HTML to PDF using reportlab while preserving formatting
                    # Parse HTML and convert to reportlab Paragraphs with formatting
                    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
                    from reportlab.lib.colors import black
                    from reportlab.lib.styles import ParagraphStyle
                    
                    logger.info("HTML content length: %d", len(html_content))
                    
                    # Create PDF in memory
                    pdf_buffer = io.BytesIO()
                    try:
                        doc = SimpleDocTemplate(pdf_buffer, pagesize=A4, 
                                                leftMargin=0.75*inch, rightMargin=0.75*inch,
                                                topMargin=0.75*inch, bottomMargin=0.75*inch)
                        styles = getSampleStyleSheet()
                        
                        # Create custom styles for better formatting
                        normal_style = ParagraphStyle(
                            'CustomNormal',
                            parent=styles['Normal'],
                            fontSize=11,
                            leading=14,
                            spaceAfter=12,
                        )
                        heading_style = ParagraphStyle(
                            'CustomHeading',
                            parent=styles['Heading1'],
                            fontSize=14,
                            leading=18,
                            spaceAfter=12,
                            textColor=black,
                        )
                        
                        story = []
                        
                        # Parse HTML and convert to reportlab Paragraphs with full formatting preservation
                        # Use BeautifulSoup if available, otherwise use simple regex-based parsing
                        try:
                            from bs4 import BeautifulSoup
                            from reportlab.platypus import ListItem, ListFlowable, KeepTogether
                            
                            soup = BeautifulSoup(html_content, 'html.parser')
                            
                            def convert_html_to_reportlab(element, parent_tag=None):
                                """Recursively convert HTML elements to reportlab format, preserving inline formatting"""
                                if element is None:
                                    return ""
                                
                                # Handle text nodes
                                if isinstance(element, str):
                                    # Escape for reportlab XML
                                    text = element
                                    text = text.replace('&', '&amp;')
                                    text = text.replace('<', '&lt;')
                                    text = text.replace('>', '&gt;')
                                    return text
                                
                                # Handle different HTML elements
                                tag_name = element.name if hasattr(element, 'name') else None
                                
                                if tag_name is None:
                                    # Text node - already handled above
                                    return ""
                                
                                result = ""
                                
                                # Handle inline formatting tags
                                if tag_name in ['strong', 'b']:
                                    inner = ''.join(convert_html_to_reportlab(child, tag_name) for child in element.children)
                                    return f'<b>{inner}</b>'
                                elif tag_name in ['em', 'i']:
                                    inner = ''.join(convert_html_to_reportlab(child, tag_name) for child in element.children)
                                    return f'<i>{inner}</i>'
                                elif tag_name in ['u']:
                                    inner = ''.join(convert_html_to_reportlab(child, tag_name) for child in element.children)
                                    return f'<u>{inner}</u>'
                                elif tag_name == 'br':
                                    return '<br/>'
                                elif tag_name in ['span', 'div']:
                                    # Just process children
                                    return ''.join(convert_html_to_reportlab(child, parent_tag) for child in element.children)
                                else:
                                    # For other tags, just process children
                                    return ''.join(convert_html_to_reportlab(child, tag_name) for child in element.children)
                            
                            # Process document structure
                            for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'div', 'ul', 'ol', 'li']):
                                # Skip if element is inside a list (we'll handle lists separately)
                                if element.find_parent(['ul', 'ol']):
                                    continue
                                
                                # Handle lists
                                if element.name in ['ul', 'ol']:
                                    list_items = []
                                    for li in element.find_all('li', recursive=False):
                                        li_text = convert_html_to_reportlab(li)
                                        if li_text.strip():
                                            list_items.append(ListItem(Paragraph(li_text, normal_style), leftIndent=0.5*inch))
                                    
                                    if list_items:
                                        # Use correct bulletType: 'bullet' for ul, '1' for ol (numbered list)
                                        bullet_type = 'bullet' if element.name == 'ul' else '1'
                                        list_flowable = ListFlowable(list_items, bulletType=bullet_type)
                                        story.append(list_flowable)
                                        story.append(Spacer(1, 0.1*inch))
                                    continue
                                
                                # Handle paragraphs and headings
                                if element.name in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'div']:
                                    # Convert HTML content to reportlab format, preserving inline formatting
                                    para_html = convert_html_to_reportlab(element)
                                    
                                    if not para_html.strip():
                                        continue
                                    
                                    # Determine style based on tag
                                    if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                                        para_style = heading_style
                                    else:
                                        para_style = normal_style
                                    
                                    try:
                                        para = Paragraph(para_html, para_style)
                                        story.append(para)
                                        story.append(Spacer(1, 0.1*inch))
                                    except Exception as para_error:
                                        logger.warning("Error creating paragraph from HTML element, skipping: %s", str(para_error))
                                    continue
                                
                                # Handle standalone list items (shouldn't happen, but just in case)
                                if element.name == 'li':
                                    li_text = convert_html_to_reportlab(element)
                                    if li_text.strip():
                                        try:
                                            para = Paragraph(li_text, normal_style)
                                            story.append(ListItem(para, leftIndent=0.5*inch))
                                            story.append(Spacer(1, 0.05*inch))
                                        except Exception as li_error:
                                            logger.warning("Error creating list item, skipping: %s", str(li_error))
                        except ImportError:
                            # Fallback: Simple text extraction if BeautifulSoup not available
                            logger.warning("BeautifulSoup not available, using simple text extraction")
                            # Remove HTML tags but preserve line breaks
                            text_content = re.sub(r'<br\s*/?>', '\n', html_content, flags=re.IGNORECASE)
                            text_content = re.sub(r'</p>', '\n\n', text_content, flags=re.IGNORECASE)
                            text_content = re.sub(r'<[^>]+>', '', text_content)
                            text_content = html.unescape(text_content)
                            text_content = re.sub(r'\n\s*\n+', '\n\n', text_content)
                            text_content = text_content.strip()
                            
                            # Split into paragraphs
                            paragraphs = text_content.split('\n\n')
                            for para_text in paragraphs:
                                para_text = para_text.strip()
                                if para_text:
                                    # Escape for reportlab
                                    para_text_escaped = para_text.replace('&', '&amp;')
                                    para_text_escaped = para_text_escaped.replace('<', '&lt;')
                                    para_text_escaped = para_text_escaped.replace('>', '&gt;')
                                    para_text_escaped = para_text_escaped.replace('\n', '<br/>')
                                    
                                    try:
                                        para = Paragraph(para_text_escaped, normal_style)
                                        story.append(para)
                                        story.append(Spacer(1, 0.15*inch))
                                    except Exception as para_error:
                                        logger.warning("Error creating paragraph, skipping: %s", str(para_error))
                        
                        # Build PDF
                        if story:
                            doc.build(story)
                            pdf_data = pdf_buffer.getvalue()
                            if pdf_data and len(pdf_data) > 0:
                                pdf_success = True
                                logger.info("Successfully converted HTML to PDF using reportlab, size: %d bytes", len(pdf_data))
                            else:
                                logger.error("PDF buffer is empty after build")
                                pdf_success = False
                                pdf_data = None
                        else:
                            logger.error("No content to add to PDF (empty story)")
                            pdf_success = False
                            pdf_data = None
                    except Exception as build_error:
                        logger.error("Error building PDF: %s", str(build_error), exc_info=True)
                        pdf_success = False
                        pdf_data = None
                    finally:
                        pdf_buffer.close()
                except ImportError as e:
                    logger.error("mammoth or reportlab not available: %s", str(e), exc_info=True)
                    pdf_success = False
                    pdf_data = None
                except Exception as e:
                    logger.error("DOCX -> HTML -> PDF conversion failed: %s", str(e), exc_info=True)
                    pdf_success = False
                    pdf_data = None
                
                # If HTML to PDF failed, try direct DOCX to PDF conversion
                if not pdf_success:
                    logger.info("Trying direct DOCX to PDF conversion")
                    success, pdf_path, error_msg = self._try_convert_docx_to_pdf(tmp_docx_path)
                    logger.info("PDF conversion result: success=%s, pdf_path=%s, error=%s", success, pdf_path, error_msg)
                    
                    if success and pdf_path and os.path.exists(pdf_path):
                        try:
                            with open(pdf_path, 'rb') as f:
                                pdf_data = f.read()
                            os.unlink(pdf_path)
                            pdf_success = True
                            logger.info("Successfully converted DOCX to PDF directly, size: %d bytes", len(pdf_data))
                        except Exception as e:
                            logger.error("Error reading PDF file: %s", str(e), exc_info=True)
                
                if pdf_success and pdf_data and len(pdf_data) > 0:
                    # Clean up temporary DOCX file
                    os.unlink(tmp_docx_path)
                    
                    # Return base64 encoded PDF data
                    base64_data = base64.b64encode(pdf_data).decode('utf-8')
                    logger.info("Successfully created PDF from DOCX for: %s, PDF size: %d bytes, base64 length: %d", 
                              document.title, len(pdf_data), len(base64_data))
                    return {
                        'preview_data': base64_data,
                        'mime_type': 'application/pdf',
                        'filename': f"{document.title}_export.pdf"
                    }
                else:
                    logger.error("All PDF conversion methods failed for document %s. export_format was: %s, pdf_success: %s, pdf_data exists: %s, pdf_data length: %s", 
                               document.title, export_format, pdf_success, pdf_data is not None, len(pdf_data) if pdf_data else 0)
                    # If PDF was explicitly requested, raise an error instead of falling back to DOCX
                    if export_format and export_format.lower() == 'pdf':
                        error_msg = "PDF conversion failed. "
                        if not pdf_success:
                            error_msg += "Conversion methods failed. "
                        if not pdf_data or len(pdf_data) == 0:
                            error_msg += "No PDF data generated. "
                        error_msg += "Please check server logs for details."
                        # Clean up temp file before raising error
                        if os.path.exists(tmp_docx_path):
                            os.unlink(tmp_docx_path)
                        raise Exception(error_msg)
                    # Only fall back to DOCX if PDF was not explicitly requested
                    # If PDF was requested and failed, we already raised an exception above
                    logger.warning("Falling back to DOCX format (PDF was not explicitly requested)")
            
            # Return DOCX format (either requested or as fallback)
            # Note: This code is only reached if export_format is not 'pdf' or if PDF conversion failed but PDF was not explicitly requested
            with open(tmp_docx_path, 'rb') as f:
                file_data = f.read()
            
            # Clean up temporary file
            os.unlink(tmp_docx_path)
            logger.debug("Temporary DOCX file cleaned up")
            
            # Return base64 encoded DOCX data
            base64_data = base64.b64encode(file_data).decode('utf-8')
            logger.info("Successfully created DOCX for: %s", document.title)
            return {
                'preview_data': base64_data,
                'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'filename': f"{document.title}_export.docx"
            }
                
        except Exception as e:
            logger.error("Error processing DOCX preview: %s", str(e), exc_info=True)
            current_app.logger.error(f"Error processing DOCX preview: {str(e)}")
            raise Exception(f"Failed to create DOCX preview: {str(e)}")
    
    def _process_html_content(self, document: Document, content_html: str, placeholder_values: Dict[str, Any], export_format: str = 'docx') -> Dict[str, Any]:
        """Process HTML content and convert to DOCX/PDF
        
        Args:
            document: Document object
            content_html: HTML content to process
            placeholder_values: Dictionary of placeholder values
            export_format: Format to export ('pdf' or 'docx')
            
        Returns:
            Dictionary with preview info including base64 data and mime type
        """
        try:
            import re
            import docx
            from docx.shared import Pt
            
            logger.debug("Processing HTML content for document: %s", document.title)
            
            # Replace placeholders in HTML
            processed_html = content_html
            placeholder_replacements = 0
            for key, value in placeholder_values.items():
                placeholder = '{{' + key + '}}'
                if placeholder in processed_html:
                    logger.debug("Replacing placeholder '%s' in HTML", key)
                    processed_html = processed_html.replace(placeholder, str(value))
                    placeholder_replacements += 1
            
            logger.info("Replaced %d placeholders in HTML content", placeholder_replacements)
            
            # Parse HTML and convert to DOCX while preserving formatting
            # Use html.parser to parse HTML structure
            from html.parser import HTMLParser
            from html import unescape
            
            class HTMLToDocxParser(HTMLParser):
                def __init__(self, doc):
                    super().__init__()
                    self.doc = doc
                    self.current_para = None
                    self.current_run = None
                    self.in_bold = False
                    self.in_italic = False
                    self.in_underline = False
                    self.in_list = False
                    self.list_type = None  # 'ul' or 'ol'
                    self.list_level = 0
                    
                def handle_starttag(self, tag, attrs):
                    if tag == 'p':
                        # Start new paragraph
                        self.current_para = self.doc.add_paragraph()
                        self.current_run = None
                    elif tag in ['strong', 'b']:
                        self.in_bold = True
                    elif tag in ['em', 'i']:
                        self.in_italic = True
                    elif tag == 'u':
                        self.in_underline = True
                    elif tag == 'ul':
                        self.in_list = True
                        self.list_type = 'ul'
                        self.list_level += 1
                    elif tag == 'ol':
                        self.in_list = True
                        self.list_type = 'ol'
                        self.list_level += 1
                    elif tag == 'li':
                        # Start new list item paragraph
                        self.current_para = self.doc.add_paragraph()
                        self.current_run = None
                        # Set list style
                        if self.list_type == 'ul':
                            self.current_para.style = 'List Bullet'
                        elif self.list_type == 'ol':
                            self.current_para.style = 'List Number'
                    elif tag == 'br':
                        # Line break within paragraph
                        if self.current_para:
                            self.current_para.add_run().add_break()
                    
                def handle_endtag(self, tag):
                    if tag == 'p':
                        # End paragraph - keep it for spacing
                        self.current_run = None
                    elif tag in ['strong', 'b']:
                        self.in_bold = False
                    elif tag in ['em', 'i']:
                        self.in_italic = False
                    elif tag == 'u':
                        self.in_underline = False
                    elif tag in ['ul', 'ol']:
                        self.list_level = max(0, self.list_level - 1)
                        if self.list_level == 0:
                            self.in_list = False
                            self.list_type = None
                    elif tag == 'li':
                        # End of list item
                        self.current_run = None
                        # Don't reset current_para here, keep it for the list item
                    
                def handle_data(self, data):
                    # Skip empty or whitespace-only data
                    data = data.strip()
                    if not data:
                        return
                    
                    # Add text to current paragraph (create one if needed)
                    if not self.current_para:
                        self.current_para = self.doc.add_paragraph()
                    
                    # Create run with formatting
                    self.current_run = self.current_para.add_run(unescape(data))
                    
                    # Apply formatting
                    if self.in_bold:
                        self.current_run.bold = True
                    if self.in_italic:
                        self.current_run.italic = True
                    if self.in_underline:
                        self.current_run.underline = True
            
            # Create DOCX document
            doc = docx.Document()
            
            # Try to parse HTML and convert to DOCX
            try:
                if processed_html and processed_html.strip():
                    parser = HTMLToDocxParser(doc)
                    parser.feed(processed_html)
                    parser.close()  # Close the parser to handle any remaining data
                else:
                    logger.warning("Empty HTML content, adding empty paragraph")
                    doc.add_paragraph(' ')
            except Exception as e:
                logger.warning("HTML parser failed, using fallback text extraction: %s", str(e), exc_info=True)
                # Fallback: simple text extraction
                try:
                    text_content = re.sub(r'<[^>]+>', '', processed_html)
                    text_content = unescape(text_content)
                    text_content = re.sub(r'\n\s*\n', '\n\n', text_content)
                    text_content = text_content.strip()
                    if text_content:
                        doc.add_paragraph(text_content)
                    else:
                        doc.add_paragraph(' ')
                except Exception as fallback_error:
                    logger.error("Fallback text extraction also failed: %s", str(fallback_error), exc_info=True)
                    doc.add_paragraph(' ')
            
            # If no paragraphs were added, add the text content as fallback
            if len(doc.paragraphs) == 0:
                # Fallback: simple text extraction
                text_content = re.sub(r'<[^>]+>', '', processed_html)
                text_content = unescape(text_content)
                text_content = re.sub(r'\n\s*\n', '\n\n', text_content)
                text_content = text_content.strip()
                if text_content:
                    doc.add_paragraph(text_content)
                else:
                    doc.add_paragraph(' ')
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
                doc.save(tmp_docx.name)
                tmp_docx_path = tmp_docx.name
                logger.debug("Saved HTML content to DOCX temp file: %s", tmp_docx_path)
            
            # Convert to requested format
            if export_format.lower() == 'pdf':
                # Try HTML to PDF conversion first (works on all platforms)
                logger.info("Attempting to convert HTML content to PDF for: %s", document.title)
                pdf_success = False
                pdf_data = None
                
                # Try using reportlab for HTML to PDF conversion (no system dependencies)
                try:
                    from reportlab.lib.pagesizes import letter, A4
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.lib.units import inch
                    import io
                    import html
                    
                    logger.info("Using reportlab for HTML to PDF conversion")
                    
                    # Convert HTML to text and create PDF
                    # Simple HTML to text conversion
                    text_content = re.sub(r'<[^>]+>', '', processed_html)
                    text_content = html.unescape(text_content)
                    text_content = re.sub(r'\n\s*\n', '\n\n', text_content)
                    text_content = text_content.strip()
                    
                    # Create PDF in memory
                    pdf_buffer = io.BytesIO()
                    doc = SimpleDocTemplate(pdf_buffer, pagesize=A4)
                    styles = getSampleStyleSheet()
                    story = []
                    
                    # Split text into paragraphs and add to PDF
                    paragraphs = text_content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            # Create paragraph with proper styling
                            para = Paragraph(para_text.strip().replace('\n', '<br/>'), styles['Normal'])
                            story.append(para)
                            story.append(Spacer(1, 0.2*inch))
                    
                    # Build PDF
                    try:
                        doc.build(story)
                        pdf_data = pdf_buffer.getvalue()
                        if pdf_data and len(pdf_data) > 0:
                            pdf_success = True
                            logger.info("Successfully converted HTML to PDF using reportlab, size: %d bytes", len(pdf_data))
                        else:
                            logger.error("PDF buffer is empty after build")
                            pdf_success = False
                    except Exception as build_error:
                        logger.error("Error building PDF: %s", str(build_error), exc_info=True)
                        pdf_success = False
                        pdf_data = None
                    finally:
                        pdf_buffer.close()
                except ImportError:
                    logger.warning("reportlab not available, trying DOCX to PDF conversion")
                except Exception as e:
                    logger.warning("reportlab conversion failed: %s, trying DOCX to PDF conversion", str(e), exc_info=True)
                
                # If HTML to PDF failed, try DOCX to PDF
                if not pdf_success:
                    logger.info("Trying DOCX to PDF conversion as fallback")
                    success, pdf_path, error_msg = self._try_convert_docx_to_pdf(tmp_docx_path)
                    
                    if success and pdf_path and os.path.exists(pdf_path):
                        try:
                            with open(pdf_path, 'rb') as f:
                                pdf_data = f.read()
                            os.unlink(pdf_path)
                            pdf_success = True
                            logger.info("Successfully converted DOCX to PDF")
                        except Exception as e:
                            logger.error("Error reading PDF file: %s", str(e))
                
                if pdf_success and pdf_data:
                    # Clean up temporary DOCX file
                    os.unlink(tmp_docx_path)
                    
                    # Return base64 encoded PDF data
                    base64_data = base64.b64encode(pdf_data).decode('utf-8')
                    logger.info("Successfully converted HTML content to PDF for: %s", document.title)
                    return {
                        'preview_data': base64_data,
                        'mime_type': 'application/pdf',
                        'filename': f"{document.title}_export.pdf"
                    }
                else:
                    logger.warning("PDF conversion failed, falling back to DOCX")
            
            # Return DOCX format (either requested or as fallback)
            with open(tmp_docx_path, 'rb') as f:
                file_data = f.read()
            
            # Clean up temp file
            os.unlink(tmp_docx_path)
            
            # Return base64 encoded DOCX data
            base64_data = base64.b64encode(file_data).decode('utf-8')
            logger.info("Successfully converted HTML content to DOCX for: %s", document.title)
            return {
                'preview_data': base64_data,
                'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'filename': f"{document.title}_export.docx"
            }
            
        except Exception as e:
            logger.error("Error processing HTML content: %s", str(e), exc_info=True)
            raise Exception(f"Failed to process HTML content: {str(e)}")
    
    def save_document_with_placeholders(self, file_data: Dict[str, Any], placeholders: List[Dict[str, Any]]) -> Document:
        """Save a document template with its placeholders
        
        Args:
            file_data: Dictionary with file information (name, description, file_path, etc.)
            placeholders: List of placeholder dictionaries
            
        Returns:
            The created Document object
        """
        try:
            logger.info("Saving document with placeholders: %s", file_data.get('name'))
            logger.debug("Document data: %s", file_data)
            logger.debug("Placeholders count: %d", len(placeholders))
            
            # Create document without requiring client and tax advisor
            document = Document(
                title=file_data.get('name'),  # Map name to title
                content=file_data.get('description', ''),
                document_type=file_data.get('file_type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),  # Use file_type as document_type
                status='active',
                file_path=file_data.get('file_path'),
                client_id=None,  # No default client required
                tax_advisor_id=None,  # No default tax advisor required
                placeholders=placeholders,  # Store placeholders directly in JSON field
                is_template=file_data.get('is_template', False),  # Mark as template if specified
                linked_client_group_ids=file_data.get('linked_client_group_ids', [])  # Store linked client group IDs
            )
            db.session.add(document)
            
            # We don't need to create separate Placeholder records anymore
            # as we're storing the placeholder data directly in the Document's JSON field
            
            db.session.commit()
            logger.info("Successfully saved document with ID %d and %d placeholders", document.id, len(placeholders))
            return document
        except Exception as e:
            logger.error("Failed to save document with placeholders: %s", str(e), exc_info=True)
            db.session.rollback()
            raise Exception(f"Failed to save document with placeholders: {str(e)}")

    def create_preview_from_uploaded_file(self, file: BinaryIO, placeholder_values: Dict[str, Any]) -> Dict[str, Any]:
        """Create a preview from an uploaded file without saving to database
        
        Args:
            file: The uploaded file object (BytesIO or similar)
            placeholder_values: Dictionary of placeholder values keyed by placeholder name
            
        Returns:
            Dictionary with preview info including base64 data and mime type
        """
        try:
            filename = getattr(file, 'filename', 'document')
            logger.info("Creating preview from uploaded file: %s", filename)
            logger.debug("Placeholder values count: %d", len(placeholder_values))
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False) as tmp:
                file.seek(0)  # Reset file pointer to beginning
                tmp.write(file.read())
                tmp_path = tmp.name
                logger.debug("Saved to temporary file: %s", tmp_path)
            
            # Determine file type
            content_type = None
            
            if filename.lower().endswith('.pdf'):
                content_type = 'application/pdf'
            elif filename.lower().endswith('.docx'):
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            else:
                logger.error("Unsupported file type: %s", filename)
                os.unlink(tmp_path)
                raise ValueError("Unsupported file type")
            
            logger.debug("Determined content type: %s", content_type)
            
            # Process based on file type
            if content_type == 'application/pdf':
                logger.info("Processing temporary PDF file")
                result = self._process_temp_pdf_preview(tmp_path, placeholder_values, filename)
            elif content_type.endswith('document'):
                logger.info("Processing temporary DOCX file")
                result = self._process_temp_docx_preview(tmp_path, placeholder_values, filename)
            else:
                logger.error("Unsupported content type: %s", content_type)
                raise ValueError(f"Unsupported content type: {content_type}")
                
            # Clean up temporary file
            os.unlink(tmp_path)
            logger.debug("Temporary file cleaned up")
            
            logger.info("Successfully created preview for uploaded file: %s", filename)
            return result
        except Exception as e:
            logger.error("Error processing temporary file preview: %s", str(e), exc_info=True)
            current_app.logger.error(f"Error processing temporary file preview: {str(e)}")
            raise Exception(f"Failed to create preview: {str(e)}")
    
    def _process_temp_pdf_preview(self, file_path: str, placeholder_values: Dict[str, Any], filename: str) -> Dict[str, Any]:
        """Process a temporary PDF file for preview
        
        Note: In a real implementation, you would use PDF manipulation libraries here
        """
        try:
            logger.debug("Processing temporary PDF preview: %s", file_path)
            # In a real implementation, you would manipulate the PDF with placeholders here
            # For now, we'll just read the file
            with open(file_path, 'rb') as f:
                file_data = f.read()
            
            # Return base64 encoded data
            base64_data = base64.b64encode(file_data).decode('utf-8')
            logger.info("Successfully created PDF preview from temporary file")
            return {
                'preview_data': base64_data,
                'mime_type': 'application/pdf',
                'filename': f"{os.path.splitext(filename)[0]}_preview.pdf"
            }
        except Exception as e:
            logger.error("Failed to process temporary PDF preview: %s", str(e), exc_info=True)
            raise Exception(f"Failed to process PDF preview: {str(e)}")
    
    def _process_temp_docx_preview(self, file_path: str, placeholder_values: Dict[str, Any], filename: str) -> Dict[str, Any]:
        """Process a temporary DOCX file for preview with placeholders
        
        First replaces placeholders in the document, then converts to PDF
        """
        try:
            import docx
            
            logger.debug("Processing temporary DOCX preview: %s", file_path)
            
            # Open the template document
            logger.debug("Opening DOCX file")
            doc = docx.Document(file_path)
            
            # Replace placeholders in all paragraphs
            placeholder_replacements = 0
            for paragraph in doc.paragraphs:
                for key, value in placeholder_values.items():
                    placeholder = '{{' + key + '}}'
                    if placeholder in paragraph.text:
                        logger.debug("Replacing placeholder '%s' with value '%s'", key, value)
                        paragraph.text = paragraph.text.replace(placeholder, str(value))
                        placeholder_replacements += 1
            
            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in placeholder_values.items():
                            placeholder = '{{' + key + '}}'
                            if placeholder in cell.text:
                                logger.debug("Replacing placeholder '%s' in table cell", key)
                                cell.text = cell.text.replace(placeholder, str(value))
                                placeholder_replacements += 1
            
            logger.info("Replaced %d placeholders in temporary DOCX", placeholder_replacements)
            
            # Save to temporary file with placeholders replaced
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
                doc.save(tmp_docx.name)
                tmp_docx_path = tmp_docx.name
                logger.debug("Saved processed DOCX to temp file: %s", tmp_docx_path)
            
            # Convert DOCX to PDF
            success, pdf_path, error_msg = self._try_convert_docx_to_pdf(tmp_docx_path)
            
            if success:
                # PDF conversion succeeded, read the file
                logger.debug("PDF created at: %s", pdf_path)
                
                try:
                    with open(pdf_path, 'rb') as f:
                        file_data = f.read()
                    
                    # Clean up temporary files
                    os.unlink(tmp_docx_path)
                    os.unlink(pdf_path)
                    logger.debug("Temporary files cleaned up")
                    
                    # Return base64 encoded PDF data
                    base64_data = base64.b64encode(file_data).decode('utf-8')
                    logger.info("Successfully created PDF preview from temporary DOCX")
                    return {
                        'preview_data': base64_data,
                        'mime_type': 'application/pdf',
                        'filename': f"{os.path.splitext(filename)[0]}_preview.pdf"
                    }
                except Exception as e:
                    logger.error("Error reading created PDF file: %s", str(e))
                    # Continue to fallback
            
            # If we got here, either conversion failed or reading the PDF failed
            # Fall back to returning the DOCX
            logger.warning("PDF conversion failed or unavailable: %s. Falling back to DOCX format", error_msg)
            
            # Read the DOCX file
            with open(tmp_docx_path, 'rb') as f:
                file_data = f.read()
            
            # Clean up temporary file
            os.unlink(tmp_docx_path)
            logger.debug("Temporary DOCX file cleaned up")
            
            # Return base64 encoded DOCX data
            base64_data = base64.b64encode(file_data).decode('utf-8')
            logger.info("Successfully created DOCX preview as fallback from temporary file")
            return {
                'preview_data': base64_data,
                'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'filename': f"{os.path.splitext(filename)[0]}_preview.docx"
            }
                
        except Exception as e:
            logger.error("Failed to process temporary DOCX preview: %s", str(e), exc_info=True)
            raise Exception(f"Failed to process DOCX preview: {str(e)}")

    def update_placeholder_name(self, document_id: int, old_name: str, new_name: str) -> bool:
        """Update a placeholder name in a document.
        
        This updates both the database record and the actual placeholders in the document file.
        
        Args:
            document_id: The ID of the document
            old_name: Current placeholder name
            new_name: New placeholder name
            
        Returns:
            Boolean indicating success
        """
        try:
            logger.info("Updating placeholder name from '%s' to '%s' in document %d", old_name, new_name, document_id)
            
            # 1. Find the document
            document = Document.query.get(document_id)
            if not document:
                logger.error("Document with ID %d not found", document_id)
                raise ValueError(f"Document with ID {document_id} not found")
                
            # 2. Update placeholder in document.placeholders JSON field
            placeholders = document.placeholders or []
            placeholder_updated = False
            
            for placeholder in placeholders:
                if placeholder.get('name') == old_name:
                    placeholder['name'] = new_name
                    placeholder_updated = True
                    logger.debug("Updated placeholder name in JSON from '%s' to '%s'", old_name, new_name)
                    
            if placeholder_updated:
                # Save the updated placeholders JSON
                document.placeholders = placeholders
                db.session.commit()
                logger.debug("Saved updated placeholders to database")
            else:
                # No placeholder in JSON, but we'll still update the document content
                logger.warning("No placeholder record found for '%s' in document %d JSON", old_name, document_id)
            
            # 3. Update the document file content based on its type
            if document.document_type.endswith('document'):  # DOCX file
                logger.info("Updating placeholder in DOCX document")
                self._update_placeholder_in_docx(document.file_path, old_name, new_name)
            elif document.document_type == 'application/pdf':
                logger.info("Updating placeholder in PDF document")
                # PDF editing is more complex and might require more specialized handling
                # For now, we'll just log this
                logger.warning("PDF placeholder update not fully implemented")
            
            logger.info("Successfully updated placeholder from '%s' to '%s' in document %d", 
                        old_name, new_name, document_id)
            return True
            
        except Exception as e:
            logger.error("Error updating placeholder name: %s", str(e), exc_info=True)
            db.session.rollback()
            raise Exception(f"Failed to update placeholder: {str(e)}")
    
    def _update_placeholder_in_docx(self, file_path: str, old_name: str, new_name: str) -> None:
        """Update placeholder names in a DOCX document.
        
        Args:
            file_path: Path to the DOCX file
            old_name: Current placeholder name
            new_name: New placeholder name
        """
        try:
            import docx
            
            logger.debug("Opening DOCX document at %s to update placeholders", file_path)
            doc = docx.Document(file_path)
            
            # Replace in paragraphs
            replacements = 0
            old_placeholder = '{{' + old_name + '}}'
            new_placeholder = '{{' + new_name + '}}'
            
            for paragraph in doc.paragraphs:
                if old_placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_placeholder, new_placeholder)
                    replacements += 1
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if old_placeholder in cell.text:
                            cell.text = cell.text.replace(old_placeholder, new_placeholder)
                            replacements += 1
            
            # Save the document if changes were made
            if replacements > 0:
                logger.info("Made %d placeholder replacements in document, saving changes", replacements)
                doc.save(file_path)
            else:
                logger.warning("No placeholders found to replace in document")
                
        except Exception as e:
            logger.error("Error updating placeholders in DOCX: %s", str(e), exc_info=True)
            raise Exception(f"Error updating DOCX placeholders: {str(e)}")

    def update_placeholder_in_temp_file(self, file: BinaryIO, old_name: str, new_name: str) -> Dict[str, Any]:
        """Update a placeholder name in a temporary document file.
        
        Args:
            file: The uploaded file object (BytesIO or similar)
            old_name: Current placeholder name
            new_name: New placeholder name
            
        Returns:
            Dictionary with updated file info including base64 data and mime type
        """
        try:
            filename = getattr(file, 'filename', 'document')
            logger.info("Updating placeholder from '%s' to '%s' in temporary file: %s", old_name, new_name, filename)
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False) as tmp:
                file.seek(0)  # Reset file pointer to beginning
                tmp.write(file.read())
                tmp_path = tmp.name
                logger.debug("Saved to temporary file: %s", tmp_path)
            
            # Determine file type
            content_type = None
            
            if filename.lower().endswith('.pdf'):
                content_type = 'application/pdf'
            elif filename.lower().endswith('.docx'):
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            else:
                logger.error("Unsupported file type: %s", filename)
                os.unlink(tmp_path)
                raise ValueError("Unsupported file type")
            
            logger.debug("Determined content type: %s", content_type)
            
            # Process based on file type
            if content_type.endswith('document'):  # DOCX file
                logger.info("Updating placeholder in DOCX file")
                # Create a new temporary file for the updated document
                updated_file_path = tmp_path + '_updated'
                
                # Copy the file first
                shutil.copy2(tmp_path, updated_file_path)
                
                # Update placeholder in the copy
                self._update_placeholder_in_docx(updated_file_path, old_name, new_name)
                
                # Read updated file
                with open(updated_file_path, 'rb') as f:
                    file_data = f.read()
                
                # Clean up temporary files
                os.unlink(tmp_path)
                os.unlink(updated_file_path)
                logger.debug("Temporary files cleaned up")
                
                # Return base64 encoded DOCX data
                base64_data = base64.b64encode(file_data).decode('utf-8')
                logger.info("Successfully updated placeholder in DOCX file")
                return {
                    'preview_data': base64_data,
                    'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'filename': filename
                }
            elif content_type == 'application/pdf':
                logger.info("PDF placeholder update not fully implemented")
                # For PDF files, we might need more complex handling
                # For now, return the original file
                with open(tmp_path, 'rb') as f:
                    file_data = f.read()
                
                # Clean up temporary file
                os.unlink(tmp_path)
                logger.debug("Temporary file cleaned up")
                
                # Return base64 encoded PDF data
                base64_data = base64.b64encode(file_data).decode('utf-8')
                logger.info("Returning original PDF (placeholder update not implemented)")
                return {
                    'preview_data': base64_data,
                    'mime_type': 'application/pdf',
                    'filename': filename
                }
            else:
                logger.error("Unsupported content type: %s", content_type)
                raise ValueError(f"Unsupported content type: {content_type}")
                
        except Exception as e:
            logger.error("Error updating placeholder in temporary file: %s", str(e), exc_info=True)
            raise Exception(f"Failed to update placeholder: {str(e)}")

    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text content from a document file
        
        Args:
            file_path: Path to the document file
            
        Returns:
            String containing the text content of the document
        """
        try:
            logger.info("Extracting text from file: %s", file_path)
            
            # Check if file exists
            if not os.path.exists(file_path):
                logger.error("File does not exist: %s", file_path)
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Determine file type by extension
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.docx':
                return self._extract_text_from_docx(file_path)
            elif file_extension == '.pdf':
                return self._extract_text_from_pdf(file_path)
            else:
                # Try to read as plain text
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        return f.read()
                except UnicodeDecodeError:
                    # Try with different encoding
                    with open(file_path, 'r', encoding='latin-1') as f:
                        return f.read()
                        
        except Exception as e:
            logger.error("Error extracting text from file %s: %s", file_path, str(e), exc_info=True)
            raise Exception(f"Failed to extract text from file: {str(e)}")
    
    def extract_text_from_uploaded_file(self, file) -> str:
        """Extract text content from an uploaded file object
        
        Args:
            file: Uploaded file object (werkzeug FileStorage)
            
        Returns:
            String containing the text content of the document
        """
        logger.info("Extracting text from uploaded file: %s", getattr(file, 'filename', 'unknown'))
        
        try:
            # Save file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{getattr(file, 'filename', 'temp')}") as temp_file:
                # Reset file pointer to beginning
                file.seek(0)
                temp_file.write(file.read())
                temp_file_path = temp_file.name
                
            logger.debug("Saved uploaded file to temporary location: %s", temp_file_path)
            
            # Extract text from temporary file
            text_content = self.extract_text_from_file(temp_file_path)
            
            # Clean up temporary file
            try:
                os.unlink(temp_file_path)
                logger.debug("Cleaned up temporary file: %s", temp_file_path)
            except OSError as e:
                logger.warning("Could not delete temporary file %s: %s", temp_file_path, str(e))
            
            return text_content
            
        except Exception as e:
            logger.error("Error extracting text from uploaded file %s: %s", 
                        getattr(file, 'filename', 'unknown'), str(e), exc_info=True)
            raise Exception(f"Failed to extract text from uploaded file: {str(e)}")
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file"""
        try:
            import docx
            
            logger.debug("Extracting text from DOCX file: %s", file_path)
            doc = docx.Document(file_path)
            
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            result = '\n'.join(text_content)
            logger.info("Successfully extracted %d characters from DOCX file", len(result))
            return result
            
        except Exception as e:
            logger.error("Error extracting text from DOCX: %s", str(e), exc_info=True)
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file"""
        try:
            # Try to use PyPDF2 if available
            try:
                import PyPDF2
                
                logger.debug("Extracting text from PDF using PyPDF2: %s", file_path)
                text_content = []
                
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text.strip():
                                text_content.append(page_text)
                        except Exception as page_error:
                            logger.warning("Failed to extract text from page %d: %s", page_num, str(page_error))
                
                result = '\n'.join(text_content)
                logger.info("Successfully extracted %d characters from PDF file using PyPDF2", len(result))
                return result
                
            except ImportError:
                logger.warning("PyPDF2 not available, trying pdfplumber")
                
                # Try pdfplumber as fallback
                try:
                    import pdfplumber
                    
                    logger.debug("Extracting text from PDF using pdfplumber: %s", file_path)
                    text_content = []
                    
                    with pdfplumber.open(file_path) as pdf:
                        for page_num, page in enumerate(pdf.pages):
                            try:
                                page_text = page.extract_text()
                                if page_text and page_text.strip():
                                    text_content.append(page_text)
                            except Exception as page_error:
                                logger.warning("Failed to extract text from page %d: %s", page_num, str(page_error))
                    
                    result = '\n'.join(text_content)
                    logger.info("Successfully extracted %d characters from PDF file using pdfplumber", len(result))
                    return result
                    
                except ImportError:
                    logger.warning("Neither PyPDF2 nor pdfplumber available for PDF text extraction")
                    return "PDF-Textextraktion ist nicht verfügbar. Bitte installieren Sie PyPDF2 oder pdfplumber."
                    
        except Exception as e:
            logger.error("Error extracting text from PDF: %s", str(e), exc_info=True)
            return f"Fehler beim Extrahieren des Textes aus der PDF-Datei: {str(e)}"

    def convert_file_to_pdf(self, file: BinaryIO, filename: str, content_type: str = None) -> Tuple[str, str]:
        """Convert uploaded file to PDF format and return the PDF file path and filename
        
        Args:
            file: The uploaded file object
            filename: Original filename
            content_type: MIME type of the file
            
        Returns:
            Tuple of (pdf_file_path, pdf_filename)
        """
        logger.info("Converting file to PDF: %s (type: %s)", filename, content_type)
        
        # First validate the file
        if not self._allowed_file(filename, content_type):
            logger.error("File type not allowed for file: %s (type: %s)", filename, content_type)
            raise ValueError(f"File type not allowed. Allowed types: {', '.join(self.allowed_extensions)}")
        
        # Determine file type
        file_extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
        
        # If it's already a PDF, just save it
        if file_extension == 'pdf':
            logger.info("File is already PDF, saving directly")
            pdf_path = self.save_document(file, filename, content_type)
            return pdf_path, filename
        
        # If it's a DOCX file, convert to PDF
        elif file_extension == 'docx':
            logger.info("Converting DOCX file to PDF")
            
            # Save the original DOCX file temporarily
            upload_folder = self._get_upload_folder()
            os.makedirs(upload_folder, exist_ok=True)
            
            # Create temporary file for the original DOCX
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
                tmp_docx.write(file.read())
                tmp_docx_path = tmp_docx.name
                logger.debug("Saved temporary DOCX file: %s", tmp_docx_path)
            
            try:
                # Convert DOCX to PDF
                success, pdf_path, error_msg = self._try_convert_docx_to_pdf(tmp_docx_path)
                
                if not success:
                    logger.error("Failed to convert DOCX to PDF: %s", error_msg)
                    # Clean up temporary file
                    os.unlink(tmp_docx_path)
                    raise ValueError(f"Failed to convert DOCX to PDF: {error_msg}")
                
                # Create final PDF filename
                base_name = os.path.splitext(filename)[0]
                pdf_filename = f"{base_name}.pdf"
                
                # Create unique filename for storage
                from datetime import datetime
                timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S_%f')
                unique_pdf_filename = f"{timestamp}_{pdf_filename}"
                
                # Move the converted PDF to the upload folder
                final_pdf_path = os.path.join(upload_folder, unique_pdf_filename)
                shutil.move(pdf_path, final_pdf_path)
                
                # Clean up temporary DOCX file
                os.unlink(tmp_docx_path)
                
                logger.info("Successfully converted DOCX to PDF: %s -> %s", filename, pdf_filename)
                return final_pdf_path, pdf_filename
                
            except Exception as e:
                # Clean up temporary file on error
                if os.path.exists(tmp_docx_path):
                    os.unlink(tmp_docx_path)
                raise
        
        # For other file types (DOC), try to handle them
        elif file_extension == 'doc':
            logger.warning("DOC files are not fully supported for conversion. Consider converting to DOCX first.")
            # For now, just save the original file without conversion
            original_path = self.save_document(file, filename, content_type)
            return original_path, filename
        
        else:
            logger.error("Unsupported file type for conversion: %s", file_extension)
            raise ValueError(f"Unsupported file type for PDF conversion: {file_extension}")

# Singleton instance
document_service = DocumentService() 